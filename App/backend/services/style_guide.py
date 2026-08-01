"""
Style guide service.
Reads BeautyFit style guide .docx files and exposes them as plain text
for prompt construction in the Pro tutorial generator.
"""

from __future__ import annotations

import logging
import os
from functools import lru_cache
from typing import Dict, List

logger = logging.getLogger(__name__)

STYLE_KEYS: List[str] = [
    "sweet",
    "natural",
    "sexy",
    "androgynous",
    "elegant",
    "mature",
]

# Alias from public style id (used on frontend) to internal guide file name
STYLE_ALIASES: Dict[str, str] = {
    "powerful": "mature",
}


def _guide_path(style: str) -> str:
    base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    filename = f"{style.capitalize()}.docx"
    return os.path.join(base, "assets", "style_guides", filename)


@lru_cache(maxsize=16)
def load_style_guide_text(style: str) -> str:
    """Load the full text content of a style guide .docx file.

    Returns an empty string when the file is missing or cannot be parsed.
    The result is cached per-style because the documents are static assets.
    """
    try:
        from docx import Document  # type: ignore
    except Exception as exc:  # pragma: no cover - optional dep not installed
        logger.warning("python-docx not available: %s", exc)
        return ""

    normalized = STYLE_ALIASES.get(style.lower(), style.lower())
    path = _guide_path(normalized)
    if not os.path.exists(path):
        logger.warning("Style guide not found: %s", path)
        return ""

    try:
        doc = Document(path)
    except Exception as exc:
        logger.error("Failed to open style guide %s: %s", path, exc)
        return ""

    paragraphs: List[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            paragraphs.append(text)

    # Also extract tables (key-value pairs often live there)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n".join(paragraphs)


def get_supported_styles() -> List[str]:
    """Return the list of styles that have a style guide available."""
    available: List[str] = []
    for style in STYLE_KEYS:
        if os.path.exists(_guide_path(style)):
            available.append(style)
    return available