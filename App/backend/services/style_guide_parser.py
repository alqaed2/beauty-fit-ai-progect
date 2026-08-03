"""Parses BeautyFit `.docx` style guides into a structured in-memory catalog.

The docx files live in `app/backend/assets/style_guides/`. Each file covers
one main style (Sweet, Natural, Sexy, Androgynous, Elegant, Mature) and
contains 5 sub-styles. For every sub-style we extract:
  - name                 (e.g. "Japanese Kawaii (日系可爱)")
  - tools                List[str]   ≈ 8–12 items
  - steps                List[{title, body}]  ≈ 5–9 items
  - pro_tips             List[str]   ≈ 4–5 items

Parsing happens ONCE at import time (and is cached) because the documents
are static assets. Results are exposed via `get_catalog()`.
"""

from __future__ import annotations

import logging
import os
import re
from functools import lru_cache
from typing import Dict, List, Optional, TypedDict

logger = logging.getLogger(__name__)

_FILES = ["Sweet", "Natural", "Sexy", "Androgynous", "Elegant", "Mature"]

# Map frontend style ids / variations -> internal catalog keys
STYLE_ALIASES: Dict[str, str] = {
    "powerful": "mature",
    "mature / powerful": "mature",
    "mature_powerful": "mature",
    "mature/powerful": "mature",
}


class ParsedStep(TypedDict):
    title: str
    body: str


class ParsedSubStyle(TypedDict):
    name: str
    tools: List[str]
    steps: List[ParsedStep]
    pro_tips: List[str]


class ParsedStyle(TypedDict):
    sub_styles: List[ParsedSubStyle]


def _guides_dir() -> str:
    """Find and return the style guides directory across different deployment environments."""
    env_dir = os.getenv("STYLE_GUIDES_DIR")
    if env_dir and os.path.exists(env_dir):
        return env_dir

    current_file_dir = os.path.dirname(os.path.abspath(__file__))
    candidates = [
        # Relative to current file: App/backend/services -> App/backend/assets/style_guides
        os.path.join(os.path.dirname(current_file_dir), "assets", "style_guides"),
        # Relative to project root
        os.path.join(os.getcwd(), "App", "backend", "assets", "style_guides"),
        os.path.join(os.getcwd(), "backend", "assets", "style_guides"),
        os.path.join(os.getcwd(), "assets", "style_guides"),
        # Render specific build path fallback
        "/opt/render/project/src/App/backend/assets/style_guides",
        "/opt/render/project/src/assets/style_guides",
    ]

    for cand in candidates:
        if os.path.exists(cand):
            logger.info("[style_guide_parser] Found style guides dir at: %s", cand)
            return cand

    default_path = os.path.join(os.path.dirname(current_file_dir), "assets", "style_guides")
    logger.warning("[style_guide_parser] Directory not found in candidates, defaulting to: %s", default_path)
    return default_path


def _parse_one_docx(path: str) -> List[ParsedSubStyle]:
    """Parse a single style-guide docx into a list of sub-style blocks."""
    try:
        from docx import Document  # type: ignore
    except Exception as exc:  # pragma: no cover
        logger.warning("python-docx not available: %s", exc)
        return []

    if not os.path.exists(path):
        logger.warning("Style guide not found: %s", path)
        return []

    try:
        doc = Document(path)
    except Exception as exc:
        logger.error("Failed to open style guide %s: %s", path, exc)
        return []

    # Flattened (style_name, text) in document order
    items: List[tuple] = []
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        sty = p.style.name if p.style else ""
        items.append((sty, text))

    # Sub-style names come from table rows like "Style N · NAME"
    sub_style_names: List[str] = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = (cell.text or "").strip()
                m = re.match(r"Style\s+(\d+)\s*[·\-]\s*(.+)", cell_text)
                if m:
                    sub_style_names.append(m.group(2).strip())

    segments: List[Dict] = []
    i = 0
    while i < len(items):
        _sty, t = items[i]
        if "Step-by-Step Tutorial" in t:
            seg: Dict = {"tools": [], "steps": [], "pro_tips": []}
            j = i + 1
            mode: Optional[str] = None
            pending_title: Optional[str] = None
            while j < len(items):
                sty2, t2 = items[j]
                if "Step-by-Step Tutorial" in t2:
                    break
                upper = t2.upper()
                if "TOOLS YOU" in upper and "NEED" in upper:
                    mode = "tools"
                    pending_title = None
                    j += 1
                    continue
                if upper == "TUTORIAL STEPS":
                    mode = "steps"
                    pending_title = None
                    j += 1
                    continue
                if upper == "PRO TIPS":
                    mode = "tips"
                    pending_title = None
                    j += 1
                    continue

                if mode == "tools" and sty2 == "List Paragraph":
                    seg["tools"].append(t2)
                elif mode == "steps":
                    if sty2 == "List Paragraph":
                        pending_title = t2
                    else:
                        if pending_title:
                            seg["steps"].append({"title": pending_title, "body": t2})
                            pending_title = None
                elif mode == "tips" and sty2 == "List Paragraph":
                    seg["pro_tips"].append(t2)
                j += 1
            segments.append(seg)
            i = j
            continue
        i += 1

    sub_styles: List[ParsedSubStyle] = []
    for idx, name in enumerate(sub_style_names):
        if idx < len(segments):
            seg = segments[idx]
        else:
            seg = {"tools": [], "steps": [], "pro_tips": []}
        sub_styles.append(
            {
                "name": name,
                "tools": list(seg["tools"]),
                "steps": [
                    {"title": str(s["title"]), "body": str(s["body"])}
                    for s in seg["steps"]
                ],
                "pro_tips": list(seg["pro_tips"]),
            }
        )
    return sub_styles


def _prebuilt_json_path() -> str:
    return os.path.join(_guides_dir(), "parsed_guides.json")


def _load_prebuilt_catalog() -> Optional[Dict[str, ParsedStyle]]:
    """Load the pre-parsed JSON catalog that ships in the assets folder."""
    import json

    path = _prebuilt_json_path()
    if not os.path.exists(path):
        return None
    try:
        with open(path, "r", encoding="utf-8") as fp:
            raw = json.load(fp)
    except Exception as exc:
        logger.warning("Failed to load prebuilt catalog %s: %s", path, exc)
        return None
    if not isinstance(raw, dict):
        return None

    catalog: Dict[str, ParsedStyle] = {}
    for key, value in raw.items():
        sub_styles_raw = (value or {}).get("sub_styles") or []
        sub_styles: List[ParsedSubStyle] = []
        for s in sub_styles_raw:
            if not isinstance(s, dict):
                continue
            steps_raw = s.get("steps") or []
            steps: List[ParsedStep] = [
                {"title": str(x.get("title", "")), "body": str(x.get("body", ""))}
                for x in steps_raw
                if isinstance(x, dict)
            ]
            sub_styles.append(
                {
                    "name": str(s.get("name", "")),
                    "tools": [str(t) for t in (s.get("tools") or [])],
                    "steps": steps,
                    "pro_tips": [str(t) for t in (s.get("pro_tips") or [])],
                }
            )
        catalog[str(key).strip().lower()] = {"sub_styles": sub_styles}
    return catalog


@lru_cache(maxsize=1)
def get_catalog() -> Dict[str, ParsedStyle]:
    """Return the fully parsed catalog: `{style_key: {sub_styles: [...]}}`."""
    prebuilt = _load_prebuilt_catalog()
    if prebuilt and any(v["sub_styles"] for v in prebuilt.values()):
        logger.info(
            "[style_guide_parser] loaded prebuilt catalog for %d styles: %s",
            len(prebuilt),
            {k: len(v["sub_styles"]) for k, v in prebuilt.items()},
        )
        return prebuilt

    catalog: Dict[str, ParsedStyle] = {}
    guides_dir = _guides_dir()
    for f in _FILES:
        path = os.path.join(guides_dir, f"{f}.docx")
        sub_styles = _parse_one_docx(path)
        catalog[f.lower()] = {"sub_styles": sub_styles}
    logger.info(
        "[style_guide_parser] loaded catalog from docx for %d styles: %s",
        len(catalog),
        {k: len(v["sub_styles"]) for k, v in catalog.items()},
    )
    return catalog


def get_style_entry(style: str) -> Optional[ParsedStyle]:
    """Return the parsed entry for a given frontend style id (with enhanced aliasing and normalization)."""
    if not style:
        return None

    normalized_style = style.strip().lower()
    key = STYLE_ALIASES.get(normalized_style, normalized_style)
    
    catalog = get_catalog()
    
    # Direct match
    if key in catalog:
        return catalog[key]

    # Flexible matching fallback (handles extra spaces or partial matches)
    for cat_key, cat_val in catalog.items():
        if cat_key == key or cat_key in key or key in cat_key:
            return cat_val

    return None


def find_sub_style(
    style: str, sub_style_name: str
) -> Optional[ParsedSubStyle]:
    """Case-insensitive sub-style lookup inside a main style."""
    entry = get_style_entry(style)
    if not entry or not entry.get("sub_styles"):
        return None

    target = (sub_style_name or "").strip().lower()
    if not target:
        return entry["sub_styles"][0] if entry["sub_styles"] else None

    target_prefix = target.split("(")[0].strip()
    for sub in entry["sub_styles"]:
        name_l = sub["name"].strip().lower()
        name_prefix = name_l.split("(")[0].strip()
        if name_l == target or name_prefix == target_prefix:
            return sub
        if target_prefix and (
            name_prefix.startswith(target_prefix)
            or target_prefix.startswith(name_prefix)
        ):
            return sub
            
    return entry["sub_styles"][0] if entry["sub_styles"] else None
